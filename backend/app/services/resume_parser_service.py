"""Extract normalized text from uploaded source resume files."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re
import unicodedata


class ResumeParserError(ValueError):
    """Raised when a source resume cannot be converted into usable text."""


_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_SPACE_RE = re.compile(r"[ \t]+")
_BLANK_LINES_RE = re.compile(r"\n{3,}")
_SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt"}


def parse_resume_bytes(
    data: bytes,
    *,
    filename: str,
    content_type: str | None = None,
) -> str:
    """Parse a PDF, DOCX, or text upload into normalized readable text."""
    if not data:
        raise ResumeParserError("Uploaded resume is empty.")

    suffix = Path(filename or "").suffix.casefold()
    if suffix not in _SUPPORTED_SUFFIXES:
        raise ResumeParserError("Upload a PDF, DOCX, or TXT resume.")

    if suffix == ".pdf" or content_type == "application/pdf":
        extracted = _parse_pdf(data)
    elif suffix == ".docx":
        extracted = _parse_docx(data)
    else:
        extracted = _decode_text(data)

    cleaned = normalize_resume_text(extracted)
    if len(cleaned) < 20:
        raise ResumeParserError("Resume text extraction returned too little readable text.")
    return cleaned


def normalize_resume_text(value: str) -> str:
    """Remove extraction artifacts without destroying human-readable resume text."""
    text = unicodedata.normalize("NFKC", value or "")
    text = text.replace("\ufeff", "").replace("\u200b", "").replace("\u2060", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _CONTROL_RE.sub("", text)
    lines = [_SPACE_RE.sub(" ", line).strip() for line in text.split("\n")]
    text = "\n".join(line for line in lines if line)
    return _BLANK_LINES_RE.sub("\n\n", text).strip()


def source_resume_file_type(filename: str) -> str:
    suffix = Path(filename or "").suffix.casefold().lstrip(".")
    return suffix or "unknown"


def _parse_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ResumeParserError("PDF parsing requires pypdf.") from exc

    try:
        reader = PdfReader(BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        raise ResumeParserError("PDF resume text extraction failed.") from exc


def _parse_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ResumeParserError("DOCX parsing requires python-docx.") from exc

    try:
        document = Document(BytesIO(data))
    except Exception as exc:
        raise ResumeParserError("DOCX resume text extraction failed.") from exc

    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ResumeParserError("Text resume encoding could not be decoded.")
