from __future__ import annotations

import re
import uuid
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.config import get_settings
from app.schemas.resume import BulletStatus, ResumeRecommendation
from app.services.jd_sanitization_service import (
    assert_render_text_safe,
    assert_resume_recommendation_safe,
    recommendation_to_plain_text,
)
from app.services.latex_render_service import sanitize_recommendation
from app.utils.latex_escape import normalize_unicode_for_resume_export


def export_recommendation_to_docx(recommendation: ResumeRecommendation) -> bytes:
    """
    Generate an ATS-safe DOCX from a ResumeRecommendation and return bytes.

    This is used as a fallback when LaTeX/PDF compilation fails.
    """
    document = _build_docx_document(recommendation)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_resume_docx(recommendation: ResumeRecommendation, generation_id: str) -> str:
    """Generate a simple single-column DOCX respecting section_order."""
    document = _build_docx_document(recommendation)
    settings = get_settings()
    output_dir = Path(settings.LATEX_OUTPUT_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"resume_{generation_id}_{uuid.uuid4().hex[:6]}.docx"
    output_path = output_dir / filename
    document.save(output_path)
    return str(output_path)


def _build_docx_document(recommendation: ResumeRecommendation) -> Document:
    rec = sanitize_recommendation(recommendation)
    from app.services.resume_validation_gate import (
        ResumeValidationError,
        validate_resume_for_export,
    )
    try:
        result = validate_resume_for_export(rec)
        rec = result.recommendation
    except ResumeValidationError:
        # The endpoint already validates before calling this function.
        # If we reach here with a fatal issue, re-raise so the caller
        # can surface the validation error to the user.
        raise
    assert_resume_recommendation_safe(rec)
    assert_render_text_safe(
        recommendation_to_plain_text(rec),
        artifact="recommendation_plain_text",
    )
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    section = document.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    if rec.contact.full_name:
        name = document.add_paragraph()
        name.alignment = 1
        run = name.add_run(_clean(rec.contact.full_name))
        run.bold = True
        run.font.size = Pt(16)

    contact = " | ".join(
        part for part in [
            _clean(rec.contact.email),
            _clean(rec.contact.phone),
            _clean(rec.contact.location),
            _clean(rec.contact.linkedin_url),
            _clean(rec.contact.github_url),
            _clean(rec.contact.portfolio_url),
        ] if part
    )
    if contact:
        paragraph = document.add_paragraph()
        paragraph.alignment = 1
        paragraph.add_run(contact)

    if rec.target_title:
        paragraph = document.add_paragraph()
        paragraph.alignment = 1
        paragraph.add_run(_clean(rec.target_title)).bold = True

    section_order = rec.section_order if rec.section_order else [
        "summary", "education", "technical skills", "experience",
        "projects", "certifications", "achievements"
    ]

    for section_name in section_order:
        if section_name == "summary" and rec.summary:
            _add_section(document, "Summary")
            document.add_paragraph(_clean(rec.summary))
        elif section_name in ("technical skills", "skills") and rec.skills:
            _add_section(document, "Skills")
            for group in rec.skills:
                if not group.skills:
                    continue
                paragraph = document.add_paragraph()
                paragraph.add_run(f"{group.category}: ").bold = True
                paragraph.add_run(", ".join(_clean(skill) for skill in group.skills if _clean(skill)))
        elif section_name == "experience" and rec.experience:
            _add_section(document, "Experience")
            for exp in rec.experience:
                if not exp.included or not exp.bullets:
                    continue
                _add_heading_line(document, exp.title, exp.company, _date_range(exp.start_date, exp.end_date or "Present"))
                for bullet in exp.bullets:
                    _add_bullet(document, bullet)
        elif section_name == "projects":
            _write_projects(document, rec)
        elif section_name == "education" and rec.education:
            _write_education(document, rec)
        elif section_name in ("achievements", "awards"):
            achievements = [*rec.achievements, *rec.awards]
            _write_achievements(document, achievements)
        elif section_name == "certifications" and rec.certifications:
            _add_section(document, "Certifications")
            for cert in rec.certifications:
                if cert.included and cert.name:
                    document.add_paragraph(" | ".join(_clean(part) for part in [cert.name, cert.issuing_org or "", cert.date or ""] if _clean(part)))

    for section in rec.custom_sections:
        if section.included and section.items:
            _add_section(document, section.title)
            for item in section.items:
                if _clean(item):
                    document.add_paragraph(_clean(item), style="List Bullet")
    return document


def _add_section(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(11)


def _add_heading_line(document: Document, title: str, subtitle: str, meta: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(_clean(title)).bold = True
    details = " | ".join(_clean(part) for part in [subtitle, meta] if _clean(part))
    if details:
        paragraph.add_run(f" | {details}")


def _add_bullet(document: Document, bullet) -> None:
    if bullet.status == BulletStatus.REJECTED:
        return
    text = _clean(bullet.text)
    if text:
        document.add_paragraph(text, style="List Bullet")


def _write_projects(document: Document, rec: ResumeRecommendation) -> None:
    if rec.projects:
        _add_section(document, "Projects")
        for project in rec.projects:
            if not project.included or not project.bullets:
                continue
            tech = ", ".join(_clean(technology) for technology in project.technologies if _clean(technology))
            _add_heading_line(document, project.name, tech, "")
            for bullet in project.bullets:
                _add_bullet(document, bullet)


def _write_education(document: Document, rec: ResumeRecommendation) -> None:
    if rec.education:
        _add_section(document, "Education")
        for edu in rec.education:
            if not edu.included:
                continue
            line = " | ".join(_clean(part) for part in [edu.degree, edu.field_of_study or "", edu.institution, f"CGPA/GPA: {edu.gpa}" if edu.gpa else ""] if _clean(part))
            if line:
                document.add_paragraph(line)


def _write_achievements(document: Document, achievements) -> None:
    if achievements:
        _add_section(document, "Achievements")
        for item in achievements:
            if item.included and item.title:
                line = " | ".join(_clean(part) for part in [item.title, item.issuer or "", item.date or ""] if _clean(part))
                document.add_paragraph(line)
                if item.description:
                    document.add_paragraph(_clean(item.description), style="List Bullet")


def _date_range(start_date: str | None, end_date: str | None) -> str:
    return " to ".join(_clean(part) for part in [start_date, end_date] if _clean(part))


def _clean(value: str | None) -> str:
    value = normalize_unicode_for_resume_export(value)
    return re.sub(r"\s+", " ", value or "").strip(" -*•\t\r\n")
